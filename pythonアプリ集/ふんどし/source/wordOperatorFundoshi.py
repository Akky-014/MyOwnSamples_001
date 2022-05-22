import docx
import renameFundoshi

# テンプレート2種類取得
templateDocrType1 = docx.Document(r'C:\techSamples\ふんどし\ふんどし差し込み１.docx')
templateDocrType2 = docx.Document(r'C:\techSamples\ふんどし\ふんどし差し込み２.docx')


# 通常版
def nomalTypeCreater(counter, katagaki1, katagaki2, shimei, keisho):
    # 段落取得
    para_katagaki1 = templateDocrType1.paragraphs[0]  # 肩書１の段落を取得
    para_katagaki2 = templateDocrType1.paragraphs[1]  # 肩書２の段落を取得
    para_shimei_keisho = templateDocrType1.paragraphs[3]  # 氏名の段落を取得

    # 肩書１編集
    t = para_katagaki1.text
    t = t.replace("肩書１", katagaki1)
    para_katagaki1.text = t  # 編集したテキストで差し替える
    templateDocrType1.paragraphs[0].runs[0].font.size = docx.shared.Pt(36)  # 文字サイズ変更

    # 肩書２編集
    t = para_katagaki2.text
    t = t.replace("肩書２", katagaki2)
    para_katagaki2.text = t  # 編集したテキストで差し替える
    # 肩書２は15桁を超えるなら、26に設定する。それ以外は36に
    if len(para_katagaki2.text) > 15:
        templateDocrType1.paragraphs[1].runs[0].font.size = docx.shared.Pt(26)
    else:
        templateDocrType1.paragraphs[1].runs[0].font.size = docx.shared.Pt(36)  # 文字サイズ変更
    
    # 氏名と敬称編集
    t = para_shimei_keisho.text
    t = t.replace("氏名", shimei + ' ' + keisho)
    para_shimei_keisho.text = t  # 編集したテキストで差し替える
    # 氏名と敬称は10桁以上なら大きさを43にする。それ以外は65に
    if len(para_shimei_keisho.text) >= 10:
        templateDocrType1.paragraphs[3].runs[0].font.size = docx.shared.Pt(43)
    else:
        templateDocrType1.paragraphs[3].runs[0].font.size = docx.shared.Pt(65)  # 文字サイズ変更

    # docx保存
    templateDocrType1.save(r'C:\techSamples\ふんどし\ふんどし_130\counter.docx')

    # ふんどし名変更処理呼び出し
    renameFundoshi.reName_path(counter, shimei, keisho)

    # 差し込み文書を元に戻す
    para_katagaki1.text = "肩書１"
    para_katagaki2.text = "肩書２"
    para_shimei_keisho.text = "氏名"

# 肩書1がなしの場合
def nonKatagaki1TypeCreater(counter, katagaki2, shimei, keisho):
    # 段落取得
    para_katagaki2 = templateDocrType2.paragraphs[0]  # 肩書２の段落を取得
    para_shimei_keisho = templateDocrType2.paragraphs[3]  # 氏名の段落を取得

    # 肩書２編集
    t = para_katagaki2.text
    t = t.replace("肩書２", katagaki2)
    para_katagaki2.text = t  # 編集したテキストで差し替える
    # 肩書２は15桁を超えるなら、26に設定する。それ以外は36に
    if len(para_katagaki2.text) > 15:
        templateDocrType2.paragraphs[0].runs[0].font.size = docx.shared.Pt(26)
    else:
        templateDocrType2.paragraphs[0].runs[0].font.size = docx.shared.Pt(36)  # 文字サイズ変更

    # 氏名と敬称編集
    keisho = ' ' + keisho
    t = para_shimei_keisho.text
    t = t.replace("氏名", shimei + keisho)
    para_shimei_keisho.text = t  # 編集したテキストで差し替える
    # 氏名と敬称は10桁以上なら大きさを43にする。それ以外は65に
    if len(para_shimei_keisho.text) >= 10:
        templateDocrType2.paragraphs[3].runs[0].font.size = docx.shared.Pt(43)
    else:
        templateDocrType2.paragraphs[3].runs[0].font.size = docx.shared.Pt(65)  # 文字サイズ変更

    templateDocrType2.save(r'C:\techSamples\ふんどし\ふんどし_130\counter.docx')

    # ふんどし名変更処理呼び出し
    renameFundoshi.reName_path(counter, shimei, keisho)

    para_katagaki2.text = "肩書２"
    para_shimei_keisho.text = "氏名"
